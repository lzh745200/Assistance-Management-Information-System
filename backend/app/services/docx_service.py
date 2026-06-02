"""
Word DOCX 报表导出服务

使用 python-docx 生成 Word 文档格式的报表，支持：
- 用户列表报表
- 村庄列表报表
- 项目列表报表
"""

import io
from datetime import datetime
from typing import Dict, List

from sqlalchemy.orm import Session

try:
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    HAS_DOCX = True
    DocumentType = Document
except ImportError:
    HAS_DOCX = False
    DocumentType = object
    Document = None  # 避免 NameError


class DocxReportService:
    """Word 报表生成服务"""

    def __init__(self):
        if not HAS_DOCX:
            raise ImportError("python-docx 未安装，请运行: pip install python-docx")

    def _create_document(self, title: str, subtitle: str = "") -> "DocumentType":
        """创建文档并设置标题"""
        doc = Document()

        # 标题
        heading = doc.add_heading(title, level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 副标题（生成时间等）
        if subtitle:
            p = doc.add_paragraph(subtitle)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(128, 128, 128)

        doc.add_paragraph("")  # 空行
        return doc

    def _add_table(self, doc: "DocumentType", headers: List[str], rows: List[List[str]]) -> None:
        """添加表格"""
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # 表头
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
            for paragraph in header_cells[i].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(10)

        # 数据行
        for row_data in rows:
            row_cells = table.add_row().cells
            for i, value in enumerate(row_data):
                row_cells[i].text = str(value)
                for paragraph in row_cells[i].paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.size = Pt(9)

    def _doc_to_bytes(self, doc: "DocumentType") -> bytes:
        """将文档转为字节"""
        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    def generate_users_report(self, users: List[Dict]) -> bytes:
        """生成用户列表报表"""
        now_str = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        doc = self._create_document("用户列表报表", f"生成时间：{now_str}    用户总数：{len(users)}")

        headers = ["ID", "用户名", "邮箱", "姓名", "角色", "状态"]
        rows = []
        for u in users:
            rows.append(
                [
                    str(u.get("id", "")),
                    u.get("username", ""),
                    u.get("email", "") or "-",
                    u.get("full_name", "") or "-",
                    u.get("role", "") or "-",
                    "启用" if u.get("is_active") else "禁用",
                ]
            )

        self._add_table(doc, headers, rows)

        # 页脚
        doc.add_paragraph("")
        footer = doc.add_paragraph(f"军队乡村振兴管理系统 — {datetime.now().year}")
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in footer.runs:
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(128, 128, 128)

        return self._doc_to_bytes(doc)

    def generate_villages_report(self, villages: List[Dict]) -> bytes:
        """生成帮扶村列表报表"""
        now_str = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        doc = self._create_document("帮扶村列表报表", f"生成时间：{now_str}    帮扶村总数：{len(villages)}")

        headers = ["ID", "村名", "县/市", "帮扶单位", "所属部门", "振兴梯队"]
        rows = []
        for v in villages:
            rows.append(
                [
                    str(v.get("id", "")),
                    v.get("village_name", "") or v.get("name", ""),
                    v.get("county", "") or "-",
                    v.get("support_unit", "") or "-",
                    v.get("department", "") or "-",
                    v.get("revitalization_tier", "") or "-",
                ]
            )

        self._add_table(doc, headers, rows)

        doc.add_paragraph("")
        footer = doc.add_paragraph(f"军队乡村振兴管理系统 — {datetime.now().year}")
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in footer.runs:
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(128, 128, 128)

        return self._doc_to_bytes(doc)

    def generate_projects_report(self, projects: List[Dict]) -> bytes:
        """生成项目列表报表"""
        now_str = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        doc = self._create_document("帮扶项目报表", f"生成时间：{now_str}    项目总数：{len(projects)}")

        headers = [
            "ID",
            "项目名称",
            "类型",
            "状态",
            "预算(万元)",
            "实际花费(万元)",
            "进度%",
        ]
        rows = []
        for p in projects:
            rows.append(
                [
                    str(p.get("id", "")),
                    p.get("name", ""),
                    p.get("type", "") or "-",
                    p.get("status", "") or "-",
                    str(p.get("budget", 0)),
                    str(p.get("actual_cost", 0)),
                    str(p.get("progress", 0)),
                ]
            )

        self._add_table(doc, headers, rows)

        doc.add_paragraph("")
        footer = doc.add_paragraph(f"军队乡村振兴管理系统 — {datetime.now().year}")
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in footer.runs:
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(128, 128, 128)

        return self._doc_to_bytes(doc)

    def generate_funds_report(self, funds: List[Dict]) -> bytes:
        """生成经费报表"""
        now_str = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        doc = self._create_document("帮扶经费报表", f"生成时间：{now_str}    记录总数：{len(funds)}")

        headers = ["ID", "名称", "金额(万元)", "状态", "经办人", "日期"]
        rows = []
        for f in funds:
            rows.append(
                [
                    str(f.get("id", "")),
                    f.get("name", "") or "-",
                    str(f.get("amount", 0)),
                    f.get("status", "") or "-",
                    f.get("operator", "") or "-",
                    str(f.get("date", "")) or "-",
                ]
            )

        self._add_table(doc, headers, rows)

        doc.add_paragraph("")
        footer = doc.add_paragraph(f"军队乡村振兴管理系统 — {datetime.now().year}")
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in footer.runs:
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(128, 128, 128)

        return self._doc_to_bytes(doc)


class DocxService:
    """向后兼容包装器 - 代理到 DocxReportService"""

    def __init__(self, db: Session = None):
        self.db = db
        self._service = DocxReportService()

    def generate_report(self, *args, **kwargs):
        return self._service.generate_executive_report(*args, **kwargs)

    @staticmethod
    def create(db: Session = None) -> "DocxService":
        return DocxService(db)
